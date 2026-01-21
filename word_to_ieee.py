#!/usr/bin/env python3
"""
Word to IEEE Format Converter
A utility to convert Microsoft Word documents to IEEE standard format.
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class IEEEFormatter:
    """Handles conversion of Word documents to IEEE format."""
    
    # IEEE Format Specifications
    IEEE_MARGINS = {
        'top': Inches(0.75),
        'bottom': Inches(1.0),
        'left': Inches(0.625),
        'right': Inches(0.625)
    }
    
    IEEE_FONTS = {
        'title': ('Times New Roman', 24, True),  # 24pt, Bold
        'author': ('Times New Roman', 10, False),  # 10pt, Regular
        'abstract_heading': ('Times New Roman', 10, True, True),  # 10pt, Bold, Italic
        'abstract_text': ('Times New Roman', 10, False),
        'section_heading': ('Times New Roman', 10, True),  # 10pt, Bold
        'subsection_heading': ('Times New Roman', 10, True, True),  # 10pt, Bold, Italic
        'body': ('Times New Roman', 10, False),  # 10pt, Regular
        'figure_caption': ('Times New Roman', 9, False, True),  # 9pt, Italic
        'table_caption': ('Times New Roman', 9, False, True),  # 9pt, Italic
        'reference': ('Times New Roman', 9, False),  # 9pt, Regular
        'code_block': ('Courier New', 9, False),  # 9pt, Monospace
        'code_caption': ('Times New Roman', 9, False)  # 9pt, Regular
    }
    
    IEEE_SPACING = {
        'before_paragraph': Pt(0),
        'after_paragraph': Pt(0),
        'line_spacing': 1.0  # Single spacing
    }
    
    def __init__(self, input_file, output_file=None, two_column=False):
        """
        Initialize the IEEE formatter.

        Args:
            input_file: Path to input Word document
            output_file: Path to output Word document (optional)
            two_column: Enable two-column format (default: False)
        """
        self.input_file = Path(input_file)
        if not self.input_file.exists():
            raise FileNotFoundError(f"Input file not found: {input_file}")

        if output_file:
            self.output_file = Path(output_file)
        else:
            self.output_file = self.input_file.parent / f"{self.input_file.stem}_IEEE{self.input_file.suffix}"

        self.doc = Document(str(self.input_file))
        self.two_column = two_column
    
    def apply_ieee_formatting(self):
        """Apply IEEE formatting to the entire document."""
        print("Applying IEEE formatting...")
        
        # Set page margins
        self._set_margins()
        
        # Format all paragraphs
        self._format_paragraphs()
        
        # Format tables
        self._format_tables()
        
        # Ensure proper section structure
        self._format_sections()
        
        print(f"Formatting complete. Saving to: {self.output_file}")
        self.doc.save(str(self.output_file))
        return self.output_file
    
    def _set_margins(self):
        """Set IEEE standard margins."""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = self.IEEE_MARGINS['top']
            section.bottom_margin = self.IEEE_MARGINS['bottom']
            section.left_margin = self.IEEE_MARGINS['left']
            section.right_margin = self.IEEE_MARGINS['right']
    
    def _format_paragraphs(self):
        """Format all paragraphs according to IEEE standards."""
        title_found = False
        abstract_found = False
        skip_next = set()  # Track paragraph indices to skip (already processed as part of code block)
        in_code_block = False  # Track if we're inside a code block

        for idx, paragraph in enumerate(self.doc.paragraphs):
            # Skip if already processed
            if idx in skip_next:
                continue

            # Skip empty paragraphs (unless inside code block)
            if not paragraph.text.strip() and not in_code_block:
                continue

            # Detect paragraph type and format accordingly
            text = paragraph.text.strip()

            # Check for code block markers
            if text.lower() == '<code block start>':
                in_code_block = True
                # Remove the marker paragraph
                paragraph.clear()
                continue
            elif text.lower() == '<code block end>':
                in_code_block = False
                # Remove the marker paragraph
                paragraph.clear()
                continue

            # If we're inside a code block, collect all paragraphs until end marker
            if in_code_block:
                code_paragraphs = [paragraph]
                next_idx = idx + 1

                # Look ahead to find all paragraphs until <code block end>
                while next_idx < len(self.doc.paragraphs):
                    next_para = self.doc.paragraphs[next_idx]
                    next_text = next_para.text.strip().lower()

                    if next_text == '<code block end>':
                        in_code_block = False
                        skip_next.add(next_idx)
                        # Remove the end marker
                        next_para.clear()
                        break
                    else:
                        code_paragraphs.append(next_para)
                        skip_next.add(next_idx)
                        next_idx += 1

                # Format all code paragraphs together
                self._format_code_block_group(code_paragraphs)
                continue

            # Title detection (usually first paragraph)
            if self._is_title(paragraph, idx):
                self._format_title(paragraph)
                title_found = True
            elif self._is_abstract_heading(paragraph):
                self._format_abstract_heading(paragraph)
                abstract_found = True
            elif self._is_author(paragraph, idx, title_found, abstract_found):
                self._format_author(paragraph)
            elif self._is_code_caption(paragraph):
                self._format_code_caption(paragraph)
            elif self._is_section_heading(paragraph):
                self._format_section_heading(paragraph)
            elif self._is_subsection_heading(paragraph):
                self._format_subsection_heading(paragraph)
            elif self._is_figure_caption(paragraph):
                self._format_figure_caption(paragraph)
            elif self._is_table_caption(paragraph):
                self._format_table_caption(paragraph)
            elif self._is_reference(paragraph):
                self._format_reference(paragraph)
            else:
                self._format_body_text(paragraph)
    
    def _is_title(self, paragraph, para_idx):
        """Check if paragraph is a title."""
        text = paragraph.text.strip()
        # Title is usually the first non-empty paragraph
        # It should not contain keywords associated with author info
        if len(text) > 0 and len(text) < 200:
            if para_idx <= 3:  # Within first few paragraphs
                # Not an author line if it contains these keywords
                if not any(kw in text.lower() for kw in ['@', 'university', 'chicago', 'manager', 'director', 'engineer', 'school']):
                    return True
        return False
    
    def _is_author(self, paragraph, para_idx, title_found, abstract_found):
        """Check if paragraph contains author information."""
        text = paragraph.text.lower()

        # Author section is only between title and abstract
        if not title_found or abstract_found:
            return False

        # Author section is typically right after title (within first 20 paragraphs)
        if para_idx > 20:
            return False

        # Check for typical author section keywords
        author_keywords = ['@', 'university', 'department', 'school of',
                          'chicago', 'usa', 'manager', 'director', 'engineer',
                          'northwestern', 'paylocity', 'duetto', 'marketplace', 'il,']

        # All paragraphs between title and abstract are author info
        # This includes names, emails, titles, organizations, locations
        return True
    
    def _is_abstract_heading(self, paragraph):
        """Check if paragraph is abstract heading."""
        text = paragraph.text.strip().lower()
        return text == 'abstract' or text.startswith('abstract -') or text.startswith('abstract-')
    
    def _is_section_heading(self, paragraph):
        """Check if paragraph is a section heading (I., II., III., etc.)."""
        text = paragraph.text.strip()
        # Check for Roman numerals or numbered sections
        if len(text) < 100 and (text.startswith('I.') or text.startswith('II.') or 
                                text.startswith('III.') or text.startswith('IV.') or
                                text.startswith('V.') or text.startswith('VI.') or
                                text.startswith('VII.') or text.startswith('VIII.') or
                                text.startswith('IX.') or text.startswith('X.')):
            return True
        # Check for numbered sections (1., 2., 3., etc.)
        if len(text) >= 2 and len(text) < 100 and text[0].isdigit() and '.' in text[:3]:
            return True
        return False
    
    def _is_subsection_heading(self, paragraph):
        """Check if paragraph is a subsection heading (A., B., C., etc.)."""
        text = paragraph.text.strip()
        if len(text) < 100 and len(text) > 2:
            # Check for letter subsections (A., B., C., etc.)
            if text[0].isalpha() and text[1] == '.':
                return True
        return False
    
    def _is_figure_caption(self, paragraph):
        """Check if paragraph is a figure caption."""
        text = paragraph.text.strip().lower()
        return text.startswith('figure') or text.startswith('fig.')
    
    def _is_table_caption(self, paragraph):
        """Check if paragraph is a table caption."""
        text = paragraph.text.strip().lower()
        return text.startswith('table')
    
    def _is_reference(self, paragraph):
        """Check if paragraph is a reference."""
        text = paragraph.text.strip()
        # References usually start with [number]
        return len(text) >= 2 and text.startswith('[') and text[1].isdigit()

    def _is_code_block(self, paragraph):
        """Check if paragraph is a code block."""
        text = paragraph.text.strip()

        # First, check if paragraph has monospace font (most reliable indicator)
        has_monospace = False
        for run in paragraph.runs:
            if run.font.name and 'courier' in run.font.name.lower():
                has_monospace = True
                break

        # If it has monospace font, it's likely a code block
        if has_monospace:
            return True

        # Otherwise, only detect as code if it has multiple strong code indicators
        # This prevents false positives from normal text

        # Must have at least 2 of these strong indicators
        strong_indicators = [
            'def ' in text and '():' in text,  # Python function definition
            'function ' in text and '{' in text,  # JavaScript function
            'class ' in text and ':' in text,  # Python class
            'CREATE TABLE' in text or 'INSERT INTO' in text or 'SELECT ' in text and 'FROM' in text,  # SQL
            text.count('(') >= 2 and text.count(')') >= 2 and ('=' in text or 'return' in text),  # Function calls
            '    ' in text and ('def ' in text or 'import ' in text or 'from ' in text),  # Indented Python
            text.startswith('{') or text.startswith('var ') or text.startswith('const '),  # JavaScript
        ]

        # Count how many strong indicators are present
        indicator_count = sum(1 for indicator in strong_indicators if indicator)

        # Also check line length and structure - code blocks are usually multi-line or have specific syntax
        has_newlines = '\n' in paragraph.text
        has_semicolons = text.count(';') >= 1
        has_indentation = '    ' in text or '\t' in text

        # Require at least 2 strong indicators OR 1 strong indicator with code-like structure
        if indicator_count >= 2:
            return True
        elif indicator_count >= 1 and (has_newlines or has_semicolons or has_indentation):
            return True

        return False

    def _is_code_caption(self, paragraph):
        """Check if paragraph is a code block caption."""
        text = paragraph.text.strip()
        return text.lower().startswith('code block')
    
    def _format_title(self, paragraph):
        """Format title paragraph."""
        font_name, font_size, is_bold = self.IEEE_FONTS['title']
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_font(paragraph, font_name, font_size, is_bold)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
    
    def _format_author(self, paragraph):
        """Format author paragraph."""
        font_name, font_size, is_bold = self.IEEE_FONTS['author']
        text = paragraph.text.strip()

        # Check if this is a job title/organization line (should be italic)
        if any(keyword in text.lower() for keyword in ['manager', 'director', 'engineer', 'school of', 'university', 'chicago', 'usa']):
            # Organization/title lines are italic
            self._apply_font(paragraph, font_name, font_size, False, True)
        else:
            # Name and email are regular (not bold, not italic)
            self._apply_font(paragraph, font_name, font_size, False, False)

        # Set formatting AFTER applying font to avoid conflicts
        # Author sections should be left-aligned, not centered
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)

        # Clear paragraph-level run properties that might override run formatting
        # This must be done AFTER setting alignment and spacing
        pPr = paragraph._element.get_or_add_pPr()
        rPr = pPr.find('.//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}rPr')
        if rPr is not None:
            pPr.remove(rPr)
    
    def _format_abstract_heading(self, paragraph):
        """Format abstract heading."""
        font_name, font_size, is_bold, is_italic = self.IEEE_FONTS['abstract_heading']
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._apply_font(paragraph, font_name, font_size, is_bold, is_italic)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    
    def _format_section_heading(self, paragraph):
        """Format section heading."""
        font_name, font_size, is_bold = self.IEEE_FONTS['section_heading']
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._apply_font(paragraph, font_name, font_size, is_bold)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    
    def _format_subsection_heading(self, paragraph):
        """Format subsection heading."""
        font_name, font_size, is_bold, is_italic = self.IEEE_FONTS['subsection_heading']
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._apply_font(paragraph, font_name, font_size, is_bold, is_italic)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
    
    def _format_figure_caption(self, paragraph):
        """Format figure caption."""
        font_name, font_size, is_bold, is_italic = self.IEEE_FONTS['figure_caption']
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_font(paragraph, font_name, font_size, is_bold, is_italic)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
    
    def _format_table_caption(self, paragraph):
        """Format table caption."""
        font_name, font_size, is_bold, is_italic = self.IEEE_FONTS['table_caption']
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_font(paragraph, font_name, font_size, is_bold, is_italic)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
    
    def _format_reference(self, paragraph):
        """Format reference paragraph."""
        font_name, font_size, is_bold = self.IEEE_FONTS['reference']
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._apply_font(paragraph, font_name, font_size, is_bold)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(0)
    
    def _format_body_text(self, paragraph):
        """Format regular body text."""
        font_name, font_size, is_bold = self.IEEE_FONTS['body']
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._apply_font(paragraph, font_name, font_size, is_bold)
        paragraph.paragraph_format.first_line_indent = Inches(0.25)  # First line indent
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = self.IEEE_SPACING['line_spacing']

    def _format_code_block_group(self, paragraphs):
        """Format a group of consecutive code block paragraphs as a single boxed unit."""
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn

        font_name, font_size, is_bold = self.IEEE_FONTS['code_block']

        for i, paragraph in enumerate(paragraphs):
            # Check if this paragraph contains newlines - if so, split it into separate lines
            if '\n' in paragraph.text:
                # This paragraph has multiple lines of code - need to preserve line breaks
                # Split the text by newlines and handle each line
                self._split_code_paragraph_by_lines(paragraph)

            # Apply monospace font to all paragraphs
            self._apply_font(paragraph, font_name, font_size, is_bold)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            pPr = paragraph._element.get_or_add_pPr()

            # Add borders and shading
            # First paragraph gets top border, last gets bottom, all get left/right
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is None:
                pBdr = OxmlElement('w:pBdr')
                pPr.append(pBdr)

            # Clear any existing borders
            for child in list(pBdr):
                pBdr.remove(child)

            # Add appropriate borders based on position
            if i == 0:
                # First paragraph: top, left, right
                for border_name in ['w:top', 'w:left', 'w:right']:
                    border = OxmlElement(border_name)
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '1')
                    border.set(qn('w:color'), '000000')
                    pBdr.append(border)
            elif i == len(paragraphs) - 1:
                # Last paragraph: bottom, left, right
                for border_name in ['w:bottom', 'w:left', 'w:right']:
                    border = OxmlElement(border_name)
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '1')
                    border.set(qn('w:color'), '000000')
                    pBdr.append(border)
            else:
                # Middle paragraphs: left and right only
                for border_name in ['w:left', 'w:right']:
                    border = OxmlElement(border_name)
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '1')
                    border.set(qn('w:color'), '000000')
                    pBdr.append(border)

            # Add light gray background to all paragraphs
            shd = pPr.find(qn('w:shd'))
            if shd is None:
                shd = OxmlElement('w:shd')
                pPr.append(shd)
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F0F0F0')

            # Set spacing and indentation for proper code formatting
            if i == 0:
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(0)
            elif i == len(paragraphs) - 1:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(6)
            else:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

            # Code block formatting - use single line spacing
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Pt(0)

    def _split_code_paragraph_by_lines(self, paragraph):
        """Split a paragraph containing newlines into proper line breaks using w:br elements."""
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn

        # Get the paragraph element
        p_elem = paragraph._element

        # Process each run in the paragraph
        for run in paragraph.runs:
            if '\n' in run.text:
                # This run contains newlines - need to split it
                r_elem = run._element

                # Get the run properties to preserve formatting
                rPr = r_elem.find(qn('w:rPr'))

                # Split the text by newlines
                lines = run.text.split('\n')

                # Clear the run's text
                for t_elem in r_elem.findall(qn('w:t')):
                    r_elem.remove(t_elem)

                # Add text elements with line breaks between them
                for idx, line in enumerate(lines):
                    # Add text element
                    t = OxmlElement('w:t')
                    t.set(qn('xml:space'), 'preserve')  # Preserve spaces
                    t.text = line
                    r_elem.append(t)

                    # Add line break after each line except the last
                    if idx < len(lines) - 1:
                        br = OxmlElement('w:br')
                        r_elem.append(br)

    def _format_code_caption(self, paragraph):
        """Format code block caption."""
        font_name, font_size, is_bold = self.IEEE_FONTS['code_caption']
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._apply_font(paragraph, font_name, font_size, is_bold)
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.left_indent = Inches(0.25)
    
    def _apply_font(self, paragraph, font_name, font_size, is_bold=False, is_italic=False):
        """Apply font formatting to paragraph."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn

        # Modify existing runs with direct XML manipulation for reliable formatting
        for run in paragraph.runs:
            rPr = run._element.get_or_add_rPr()

            # Set font name
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
                rPr.append(rFonts)
            else:
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)

            # Set font size (in half-points)
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = parse_xml(f'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{font_size * 2}"/>')
                rPr.append(sz)
            else:
                sz.set(qn('w:val'), str(font_size * 2))

            # Set bold
            b = rPr.find(qn('w:b'))
            if is_bold:
                if b is None:
                    b = parse_xml('<w:b xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    rPr.append(b)
                else:
                    # Remove w:val attribute if present
                    if b.get(qn('w:val')) is not None:
                        del b.attrib[qn('w:val')]
            else:
                # Remove bold element entirely or set to 0
                if b is not None:
                    rPr.remove(b)

            # Set italic
            i = rPr.find(qn('w:i'))
            if is_italic:
                if i is None:
                    i = parse_xml('<w:i xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    rPr.append(i)
                else:
                    if i.get(qn('w:val')) is not None:
                        del i.attrib[qn('w:val')]
            else:
                if i is not None:
                    rPr.remove(i)
    
    def _format_tables(self):
        """Format all tables according to IEEE standards."""
        for table in self.doc.tables:
            # Set table font
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(9)  # Table text is typically 9pt
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _format_sections(self):
        """Ensure proper section formatting."""
        if self.two_column:
            # Apply two-column format to all sections
            # Title and author sections should remain single column
            # Body content (after abstract) should be two-column
            self._apply_two_column_layout()

    def _apply_two_column_layout(self):
        """Apply two-column layout to document sections."""
        # Find where to start two-column format (after title and author info)
        # Typically starts from abstract or first section

        # Create a new section for two-column layout
        # The first section (title, authors) remains single column
        # Subsequent sections use two columns

        sections = self.doc.sections

        # If document has only one section, we need to add a section break
        # after the title/author area to enable two-column formatting
        if len(sections) == 1:
            # Add section break before abstract or first section heading
            abstract_found = False
            for i, paragraph in enumerate(self.doc.paragraphs):
                text = paragraph.text.strip().lower()
                # Start two-column from abstract or first section heading
                if text == 'abstract' or self._is_section_heading(paragraph):
                    # Insert a section break before this paragraph
                    # Note: python-docx doesn't support inserting section breaks mid-document easily
                    # So we'll apply two-column to all sections for now
                    abstract_found = True
                    break

        # Apply two-column format to all sections
        # In a real implementation, you'd want the first section (title/authors) single column
        # and remaining sections two-column
        for section in sections:
            # Two columns with 0.5 inch spacing
            section.page_width = Inches(8.5)  # US Letter width
            section.page_height = Inches(11)  # US Letter height

            # For two-column IEEE format:
            # Each column is approximately 3.5 inches wide
            # with 0.25 inches spacing between columns
            # This is handled by Word's column settings

            # Note: python-docx has limited support for column formatting
            # The _element property gives access to the underlying XML
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls

            # Create two-column layout using Word's XML structure
            cols_xml = parse_xml(
                f'<w:cols {nsdecls("w")} w:num="2" w:space="360"/>'
            )
            section._sectPr.append(cols_xml)


def main():
    """Main entry point for the converter."""
    import argparse

    parser = argparse.ArgumentParser(
        description='Convert Microsoft Word documents to IEEE standard format.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  python word_to_ieee.py paper.docx
  python word_to_ieee.py paper.docx paper_ieee.docx
  python word_to_ieee.py paper.docx --two-column
  python word_to_ieee.py paper.docx -o output.docx --two-column
        '''
    )

    parser.add_argument('input_file', help='Input Word document (.docx)')
    parser.add_argument('output_file', nargs='?', default=None,
                        help='Output Word document (optional)')
    parser.add_argument('-o', '--output', dest='output_alt', default=None,
                        help='Alternative way to specify output file')
    parser.add_argument('--two-column', '-2', action='store_true',
                        help='Enable two-column format (IEEE standard)')
    parser.add_argument('--version', action='version', version='%(prog)s 1.1.0')

    args = parser.parse_args()

    input_file = args.input_file
    output_file = args.output_file or args.output_alt
    two_column = args.two_column

    try:
        formatter = IEEEFormatter(input_file, output_file, two_column=two_column)
        output_path = formatter.apply_ieee_formatting()

        if two_column:
            print(f"\n[SUCCESS] IEEE-formatted document (two-column) saved to: {output_path}")
        else:
            print(f"\n[SUCCESS] IEEE-formatted document saved to: {output_path}")
            print("         Use --two-column flag for two-column format")
    except Exception as e:
        print(f"\n[ERROR] {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
