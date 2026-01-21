#!/usr/bin/env python3
"""Create a test document with code blocks for testing IEEE formatting."""

from docx import Document
from docx.shared import Pt

# Create a test document
doc = Document()

# Add title
title = doc.add_paragraph('A Test Paper with Code Examples')
title.alignment = 1  # Center

# Add authors
doc.add_paragraph('John Doe')
doc.add_paragraph('john.doe@university.edu')
doc.add_paragraph('Department of Computer Science')
doc.add_paragraph('University of Example')
doc.add_paragraph('Chicago, IL, USA')

# Add abstract
doc.add_paragraph('Abstract')
doc.add_paragraph('This is a test document to demonstrate code block formatting in IEEE style. The document includes various code examples in different programming languages.')

# Add a section
doc.add_paragraph('I. INTRODUCTION')
doc.add_paragraph('This document demonstrates automatic code block detection and formatting.')

# Add section II
doc.add_paragraph('II. CODE EXAMPLES')

# Add subsection
doc.add_paragraph('A. Python Example')
doc.add_paragraph('The following shows a Python function:')

# Add code block start marker
doc.add_paragraph('<code block start>')

# Add Python code block
code_para1 = doc.add_paragraph()
code_run1 = code_para1.add_run(
    'def fetch_trait_profile(profile_id):\n'
    '    """Retrieve trait profile from database."""\n'
    '    result = db.query(\n'
    '        "SELECT * FROM TraitProfiles WHERE Profile_ID=?",\n'
    '        profile_id\n'
    '    )\n'
    '    return result'
)
code_run1.font.name = 'Courier New'
code_run1.font.size = Pt(9)

# Add code block end marker
doc.add_paragraph('<code block end>')

# Add code caption
doc.add_paragraph('Code Block 1: Python Function Example')

# Add subsection
doc.add_paragraph('B. SQL Example')
doc.add_paragraph('Database schema definition:')

# Add code block start marker
doc.add_paragraph('<code block start>')

# Add SQL code block
code_para2 = doc.add_paragraph()
code_run2 = code_para2.add_run(
    'CREATE TABLE TraitProfiles (\n'
    '    Profile_ID VARCHAR(50) PRIMARY KEY,\n'
    '    Version VARCHAR(10) NOT NULL,\n'
    '    Domain VARCHAR(30) NOT NULL,\n'
    '    Traits JSON NOT NULL\n'
    ');'
)
code_run2.font.name = 'Courier New'
code_run2.font.size = Pt(9)

# Add code block end marker
doc.add_paragraph('<code block end>')

# Add code caption
doc.add_paragraph('Code Block 2: SQL Schema Definition')

# Add subsection
doc.add_paragraph('C. JavaScript Example')
doc.add_paragraph('An example JavaScript function:')

# Add code block start marker
doc.add_paragraph('<code block start>')

# Add JavaScript code block
code_para3 = doc.add_paragraph()
code_run3 = code_para3.add_run(
    'function generateEmbedding(traits) {\n'
    '    const weights = {\n'
    '        empathy: 0.2,\n'
    '        humor: 0.15,\n'
    '        ethics: 0.25\n'
    '    };\n'
    '    return weights;\n'
    '}'
)
code_run3.font.name = 'Courier New'
code_run3.font.size = Pt(9)

# Add code block end marker
doc.add_paragraph('<code block end>')

# Add code caption
doc.add_paragraph('Code Block 3: JavaScript Function')

# Add conclusion section
doc.add_paragraph('III. CONCLUSION')
doc.add_paragraph('This document successfully demonstrates code block formatting.')

# Add references
doc.add_paragraph('[1] J. Smith, "Sample Reference," IEEE Journal, 2024.')

# Save document
output_path = 'examples/test_code_blocks.docx'
doc.save(output_path)
print(f'Test document created: {output_path}')
