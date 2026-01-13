#!/usr/bin/env python3
"""Create a sample Word document for testing the IEEE converter."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_document():
    """Create a sample academic paper for testing."""
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("A Sample Research Paper for Testing IEEE Format Conversion")
    title_run.font.size = Pt(16)
    title_run.bold = True
    
    # Author
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run("John Doe, Jane Smith")
    author.runs[0].font.size = Pt(12)
    
    # Abstract heading
    doc.add_paragraph("Abstract")
    abstract_para = doc.add_paragraph()
    abstract_para.add_run("This is a sample abstract for testing the IEEE format converter. The converter should properly format this document according to IEEE publication standards, including margins, fonts, spacing, and section headings.")
    
    # Section I
    doc.add_paragraph("I. Introduction")
    intro = doc.add_paragraph()
    intro.add_run("This paper presents a utility for converting Microsoft Word documents to IEEE standard format. The conversion process involves applying proper formatting, fonts, margins, and styles according to IEEE publication guidelines.")
    
    # Section II
    doc.add_paragraph("II. Methodology")
    method = doc.add_paragraph()
    method.add_run("The methodology involves automatic detection of document elements such as titles, sections, subsections, figures, and tables. Each element is then formatted according to IEEE specifications.")
    
    # Subsection
    doc.add_paragraph("A. Data Collection")
    data = doc.add_paragraph()
    data.add_run("Data was collected from various sources to validate the formatting requirements.")
    
    # Section III
    doc.add_paragraph("III. Results")
    results = doc.add_paragraph()
    results.add_run("The results demonstrate successful conversion of documents to IEEE format with proper formatting applied to all elements.")
    
    # Figure caption
    fig_caption = doc.add_paragraph()
    fig_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_caption.add_run("Figure 1. Sample figure caption for testing.")
    
    # Section IV
    doc.add_paragraph("IV. Conclusion")
    conclusion = doc.add_paragraph()
    conclusion.add_run("In conclusion, the utility successfully converts Word documents to IEEE format, making it easier for researchers to prepare their papers for publication.")
    
    # References
    doc.add_paragraph("References")
    ref1 = doc.add_paragraph("[1] J. Smith, \"Sample Reference,\" IEEE Transactions, vol. 1, no. 1, pp. 1-10, 2024.")
    ref2 = doc.add_paragraph("[2] A. Johnson, \"Another Reference,\" in Proc. Conference, 2024, pp. 20-30.")
    
    # Save
    output_file = "test_document.docx"
    doc.save(output_file)
    print(f"Test document created: {output_file}")
    return output_file

if __name__ == "__main__":
    create_test_document()
